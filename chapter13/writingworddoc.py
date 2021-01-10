#! /usr/bin/env python3

import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph')
paraObj2 = doc.add_paragraph('This is yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph')

# Optional 2nd parameter for styling and we can also add headings
doc.add_paragraph('World Hello!', 'Title')
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)

# Adding line and page breaks
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')

# Adding Pictures
doc.add_picture('zophie.png', width = docx.shared.Inches(1),
        height = docx.shared.Cm(4))

doc.save('helloworld.docx')
